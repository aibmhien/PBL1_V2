from __future__ import annotations

import os
import sqlite3
import hashlib
import secrets
import random
import zipfile
import shutil
import io
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from fastapi import FastAPI, File, Form, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, PlainTextResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from openpyxl import load_workbook
import qrcode
try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except Exception:
    HTML = None
    WEASYPRINT_AVAILABLE = False

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = Path(os.environ.get("PBL_DATA_DIR", BASE_DIR / "data"))
GENERATED_DIR = DATA_DIR / "generated"
SUBMISSIONS_DIR = DATA_DIR / "submissions"
MATERIALS_DIR = DATA_DIR / "materials"
BACKUPS_DIR = DATA_DIR / "backups"
DEMO_DIR = BASE_DIR / "demo"
DB_PATH = DATA_DIR / "pbl.db"
STATIC_DIR = BASE_DIR / "static"
TEMPLATES_DIR = BASE_DIR / "templates"

# Render/Git deployments may omit an empty directory. Create required runtime
# directories before FastAPI/Starlette mounts them so startup never fails.
STATIC_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)
GENERATED_DIR.mkdir(exist_ok=True)
SUBMISSIONS_DIR.mkdir(exist_ok=True)
MATERIALS_DIR.mkdir(exist_ok=True)
BACKUPS_DIR.mkdir(exist_ok=True)

app = FastAPI(title="PBL Manager V0.4.2")
ADMIN_PASSWORD = os.environ.get("PBL_ADMIN_PASSWORD", "pbl123")
AUTH_SECRET = os.environ.get("PBL_AUTH_SECRET", hashlib.sha256((ADMIN_PASSWORD + "|pbl-v040").encode()).hexdigest())

def _sign(value: str) -> str:
    return hashlib.sha256((value + "|" + AUTH_SECRET).encode()).hexdigest()

def _staff_from_cookie(request: Request):
    raw = request.cookies.get("pbl_staff", "")
    if not raw or ":" not in raw:
        return None
    kind, ident, sig = (raw.split(":", 2) + ["", ""])[:3] if raw.count(":") >= 2 else ("", "", "")
    if not secrets.compare_digest(sig, _sign(f"{kind}:{ident}")):
        return None
    if kind == "admin":
        return {"role": "admin", "name": get_settings().get("instructor_name", "Giảng viên hướng dẫn"), "username": "admin"}
    if kind == "reviewer":
        try:
            uid = int(ident)
        except Exception:
            return None
        with db() as c:
            row = c.execute("SELECT id,username,display_name,role,active FROM staff_users WHERE id=?", (uid,)).fetchone()
        if row and row["active"]:
            return {"role": row["role"], "name": row["display_name"], "username": row["username"], "id": row["id"]}
    return None

@app.middleware("http")
async def protect_staff(request: Request, call_next):
    path = request.url.path
    staff = _staff_from_cookie(request)
    request.state.staff = staff
    public_prefixes = ("/student", "/static", "/attendance/checkin", "/material/")
    public_exact = ("/admin-login", "/admin-logout")
    is_group_staff = path.startswith("/group/") and not path.endswith("/print") and not path.endswith("/pdf")
    staff_only = path == "/" or path.startswith("/admin") or is_group_staff or path.startswith("/submission/") or path == "/materials"
    if staff_only and path not in public_exact and not staff:
        return RedirectResponse(url="/admin-login", status_code=303)
    # Reviewer is read/grade only: cannot change assignment/templates/settings/materials/accounts/attendance sessions/backups.
    if staff and staff.get("role") == "reviewer":
        admin_allowed = (request.method == "GET" and (path in ("/admin/attendance", "/admin/plan") or path.startswith("/admin/attendance/view/")))
        grading_allowed = path.startswith("/group/") and "/progress/" in path
        if path.startswith("/admin") and not admin_allowed and not grading_allowed:
            return PlainTextResponse("Tài khoản đồng nghiệp chỉ có quyền xem, theo dõi và cho điểm/nhận xét.", status_code=403)
    return await call_next(request)
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
templates = Jinja2Templates(directory=TEMPLATES_DIR)

PROJECT_INFO = {
    1: {"title": "Dự án số 1", "task": "Thiết kế hệ thống cơ khí sử dụng trong vận chuyển hàng hóa.", "load": "Theo biểu đồ đặc tính tải trọng", "direction": "Làm việc một chiều"},
    2: {"title": "Dự án số 2", "task": "Thiết kế hệ thống cơ khí sử dụng để nâng hạ hàng hóa.", "load": "Tải thay đổi, rung động vừa", "direction": "Làm việc hai chiều"},
    3: {"title": "Dự án số 3", "task": "Thiết kế hệ thống cơ khí sử dụng trong vận chuyển hàng hóa.", "load": "Tải trọng thay đổi, rung động", "direction": "Làm việc một chiều"},
    4: {"title": "Dự án số 4", "task": "Thiết kế hệ thống cơ khí sử dụng để nâng hạ hàng hóa.", "load": "Tải trọng thay đổi, rung động nhẹ", "direction": "Làm việc hai chiều"},
    5: {"title": "Dự án số 5", "task": "Thiết kế hệ thống cơ khí sử dụng trong vận chuyển hàng hóa.", "load": "Tải trọng thay đổi, rung động nhẹ", "direction": "Làm việc một chiều"},
}

DEFAULT_MILESTONES = [
    (1, "Nhận đề & phân tích nhiệm vụ"),
    (2, "Tính chọn động cơ / phân phối tỉ số truyền"),
    (3, "Thiết kế các bộ truyền và chi tiết chính"),
    (4, "Hoàn thiện bản vẽ & báo cáo"),
]


def hash_password(password: str, salt: str | None = None) -> str:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 120000).hex()
    return f"{salt}${digest}"

def verify_password(password: str, stored: str) -> bool:
    try:
        salt, digest = stored.split("$", 1)
        return secrets.compare_digest(hash_password(password, salt).split("$", 1)[1], digest)
    except Exception:
        return False

def current_staff(request: Request):
    return getattr(request.state, "staff", None)

def db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_db():
    with db() as c:
        c.executescript(
            """
            CREATE TABLE IF NOT EXISTS students (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                stt INTEGER,
                student_id TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                project_type INTEGER NOT NULL
            );
            CREATE TABLE IF NOT EXISTS project_data (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_type INTEGER NOT NULL,
                data_code TEXT NOT NULL,
                p REAL, v REAL, d REAL, t REAL, nam REAL, ngay REAL,
                UNIQUE(project_type, data_code)
            );
            CREATE TABLE IF NOT EXISTS groups_tbl (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_code TEXT UNIQUE NOT NULL,
                project_type INTEGER NOT NULL,
                data_id INTEGER,
                FOREIGN KEY(data_id) REFERENCES project_data(id)
            );
            CREATE TABLE IF NOT EXISTS group_students (
                group_id INTEGER NOT NULL,
                student_id INTEGER NOT NULL,
                PRIMARY KEY(group_id, student_id),
                FOREIGN KEY(group_id) REFERENCES groups_tbl(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS templates (
                project_type INTEGER PRIMARY KEY,
                file_path TEXT NOT NULL,
                version INTEGER NOT NULL DEFAULT 1,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS milestones (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                position INTEGER NOT NULL,
                title TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS group_progress (
                group_id INTEGER NOT NULL,
                milestone_id INTEGER NOT NULL,
                status TEXT NOT NULL DEFAULT 'Chưa bắt đầu',
                score REAL,
                feedback TEXT,
                PRIMARY KEY(group_id, milestone_id),
                FOREIGN KEY(group_id) REFERENCES groups_tbl(id) ON DELETE CASCADE,
                FOREIGN KEY(milestone_id) REFERENCES milestones(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS app_settings (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL DEFAULT ''
            );
            CREATE TABLE IF NOT EXISTS submissions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_id INTEGER NOT NULL,
                student_id INTEGER NOT NULL,
                milestone_id INTEGER,
                note TEXT,
                original_name TEXT NOT NULL,
                stored_name TEXT NOT NULL,
                file_type TEXT NOT NULL,
                uploaded_at TEXT NOT NULL,
                FOREIGN KEY(group_id) REFERENCES groups_tbl(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                FOREIGN KEY(milestone_id) REFERENCES milestones(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS staff_users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                display_name TEXT NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL DEFAULT 'reviewer',
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS materials (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                description TEXT,
                kind TEXT NOT NULL,
                url TEXT,
                original_name TEXT,
                stored_name TEXT,
                created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS attendance_sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                session_date TEXT NOT NULL,
                token TEXT UNIQUE NOT NULL,
                expires_at TEXT NOT NULL,
                is_open INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS attendance_records (
                session_id INTEGER NOT NULL,
                student_id INTEGER NOT NULL,
                status TEXT NOT NULL DEFAULT 'Có mặt',
                checked_at TEXT NOT NULL,
                PRIMARY KEY(session_id, student_id),
                FOREIGN KEY(session_id) REFERENCES attendance_sessions(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS schema_meta (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS evaluation_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_id INTEGER NOT NULL,
                milestone_id INTEGER NOT NULL,
                evaluator TEXT NOT NULL,
                score REAL,
                feedback TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(group_id) REFERENCES groups_tbl(id) ON DELETE CASCADE,
                FOREIGN KEY(milestone_id) REFERENCES milestones(id) ON DELETE CASCADE
            );
            """
        )
        # Migration bổ sung cho kế hoạch chung toàn lớp; giữ nguyên dữ liệu cũ.
        milestone_cols = {r["name"] for r in c.execute("PRAGMA table_info(milestones)").fetchall()}
        if "description" not in milestone_cols:
            c.execute("ALTER TABLE milestones ADD COLUMN description TEXT NOT NULL DEFAULT ''")
        if "start_at" not in milestone_cols:
            c.execute("ALTER TABLE milestones ADD COLUMN start_at TEXT NOT NULL DEFAULT ''")
        if "due_at" not in milestone_cols:
            c.execute("ALTER TABLE milestones ADD COLUMN due_at TEXT NOT NULL DEFAULT ''")
        if "is_active" not in milestone_cols:
            c.execute("ALTER TABLE milestones ADD COLUMN is_active INTEGER NOT NULL DEFAULT 0")

        if c.execute("SELECT COUNT(*) n FROM milestones").fetchone()["n"] == 0:
            c.executemany("INSERT INTO milestones(position,title) VALUES (?,?)", DEFAULT_MILESTONES)
        defaults = {
            "instructor_name": "TS. Bùi Minh Hiển",
            "start_week": "1",
            "end_week": "15",
            "academic_year_start": "2026",
            "academic_year_end": "2027",
            "attendance_weight": "10",
        }
        for key, value in defaults.items():
            c.execute("INSERT OR IGNORE INTO app_settings(key,value) VALUES (?,?)", (key, value))
        c.execute("INSERT INTO schema_meta(key,value) VALUES ('schema_version','0.4.1') ON CONFLICT(key) DO UPDATE SET value=excluded.value")


def normalize_project_type(value):
    """Chấp nhận mã đề I-V hoặc 1-5; để trống trả về 0."""
    if value in (None, ""):
        return 0
    text = str(value).strip().upper()
    roman = {"I":1, "II":2, "III":3, "IV":4, "V":5, "DA1":1, "DA2":2, "DA3":3, "DA4":4, "DA5":5}
    if text in roman:
        return roman[text]
    try:
        n = int(float(text))
    except Exception as exc:
        raise ValueError(f"Mã Đề={value} không hợp lệ. Dùng I-V hoặc 1-5.") from exc
    if n not in PROJECT_INFO:
        raise ValueError(f"Mã Đề={value} chưa được hỗ trợ. Dùng I-V hoặc 1-5.")
    return n


def parse_students_xlsx(path: Path):
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    header_row = None
    cols = {}
    for r in range(1, min(ws.max_row, 30) + 1):
        norm = [str(ws.cell(r, c).value or "").strip().lower() for c in range(1, ws.max_column + 1)]
        if any(v in ("số thẻ", "mssv", "mã sv", "ma sv") for v in norm) and any(v in ("họ tên", "họ và tên", "ho ten") for v in norm):
            header_row = r
            for i, v in enumerate(norm, 1):
                if v in ("tt", "stt"):
                    cols["stt"] = i
                elif v in ("số thẻ", "mssv", "mã sv", "ma sv"):
                    cols["student_id"] = i
                elif v in ("họ tên", "họ và tên", "ho ten"):
                    cols["name"] = i
                elif v in ("đề", "de", "mã đề", "ma de"):
                    cols["project_type"] = i
            break
    if not header_row or not all(k in cols for k in ("student_id", "name")):
        raise ValueError("Không tìm thấy hàng tiêu đề gồm Số thẻ/MSSV và Họ tên.")
    result = []
    for r in range(header_row + 1, ws.max_row + 1):
        sid = ws.cell(r, cols["student_id"]).value
        name = ws.cell(r, cols["name"]).value
        ptype = ws.cell(r, cols["project_type"]).value if "project_type" in cols else None
        if sid is None and name is None:
            continue
        if sid is None or name is None:
            continue
        try:
            ptype = normalize_project_type(ptype)
        except ValueError as exc:
            raise ValueError(f"Dòng {r}: {exc}") from exc
        stt_raw = ws.cell(r, cols["stt"]).value if "stt" in cols else len(result) + 1
        try:
            stt = int(stt_raw)
        except Exception:
            stt = len(result) + 1
        result.append((stt, str(sid).strip(), str(name).strip(), ptype))
    return result


def parse_project_data(path: Path, max_project=5):
    out = []
    if path.suffix.lower() == ".xlsx":
        wb = load_workbook(path, data_only=True)
        for ptype in range(1, max_project + 1):
            name = f"De-{ptype}"
            if name not in wb.sheetnames:
                raise ValueError(f"Thiếu sheet {name}")
            ws = wb[name]
            for r in range(2, ws.max_row + 1):
                vals = [ws.cell(r, c).value for c in range(1, 8)]
                if vals[0] in (None, ""):
                    continue
                out.append((ptype, str(vals[0]).strip(), *vals[1:7]))
    elif path.suffix.lower() == ".xls":
        try:
            import xlrd
        except ImportError as exc:
            raise ValueError("Đọc trực tiếp .xls cần gói xlrd. Có thể Save As file sang .xlsx rồi import lại.") from exc
        wb = xlrd.open_workbook(path)
        for ptype in range(1, max_project + 1):
            name = f"De-{ptype}"
            try:
                ws = wb.sheet_by_name(name)
            except Exception as exc:
                raise ValueError(f"Thiếu sheet {name}") from exc
            for r in range(1, ws.nrows):
                vals = ws.row_values(r, 0, min(7, ws.ncols)) + [None] * max(0, 7 - ws.ncols)
                if vals[0] in (None, ""):
                    continue
                out.append((ptype, str(vals[0]).strip(), *vals[1:7]))
    else:
        raise ValueError("Chỉ hỗ trợ .xls hoặc .xlsx")
    return out


def next_group_code(c):
    rows = c.execute("SELECT group_code FROM groups_tbl").fetchall()
    nums = []
    for r in rows:
        code = str(r["group_code"] or "")
        if code.startswith("N") and code[1:].isdigit():
            nums.append(int(code[1:]))
    return f"N{(max(nums) if nums else 0)+1:02d}"


def assign_next_unused_data(c, group_id: int, project_type: int):
    row = c.execute(
        """SELECT pd.id FROM project_data pd
           WHERE pd.project_type=? AND pd.id NOT IN
           (SELECT data_id FROM groups_tbl WHERE data_id IS NOT NULL AND id<>?)
           ORDER BY pd.id LIMIT 1""",
        (project_type, group_id),
    ).fetchone()
    c.execute("UPDATE groups_tbl SET data_id=? WHERE id=?", (row["id"] if row else None, group_id))


def assign_random_unused_data(c, group_id: int, project_type: int):
    rows = c.execute(
        """SELECT pd.id FROM project_data pd WHERE pd.project_type=? AND pd.id NOT IN
        (SELECT data_id FROM groups_tbl WHERE data_id IS NOT NULL AND id<>?)""",
        (project_type, group_id),
    ).fetchall()
    if not rows:
        return False
    chosen = random.choice(rows)["id"]
    c.execute("UPDATE groups_tbl SET data_id=? WHERE id=?", (chosen, group_id))
    return True

def random_assign_groups(only_unassigned: bool = True):
    with db() as c:
        # Tính năng V0.4: nếu danh sách chưa có cột Đề / chưa có nhóm,
        # ghép các sinh viên chưa thuộc nhóm theo đúng thứ tự 1-2, 3-4,...
        ungrouped = c.execute("""SELECT s.* FROM students s WHERE NOT EXISTS
            (SELECT 1 FROM group_students gs WHERE gs.student_id=s.id) ORDER BY s.stt,s.id""").fetchall()
        for i in range(0, len(ungrouped), 2):
            members = ungrouped[i:i+2]
            code = next_group_code(c)
            cur = c.execute("INSERT INTO groups_tbl(group_code,project_type) VALUES (?,0)", (code,))
            gid = cur.lastrowid
            c.executemany("INSERT INTO group_students(group_id,student_id) VALUES (?,?)", [(gid,m["id"]) for m in members])
            ensure_progress_rows(c, gid)
        used_counts = {p: c.execute("SELECT COUNT(*) n FROM groups_tbl WHERE project_type=?", (p,)).fetchone()["n"] for p in PROJECT_INFO}
        if only_unassigned:
            groups = c.execute("SELECT * FROM groups_tbl WHERE project_type=0 OR data_id IS NULL ORDER BY id").fetchall()
        else:
            groups = c.execute("SELECT * FROM groups_tbl ORDER BY id").fetchall()
            for g in groups:
                c.execute("UPDATE groups_tbl SET project_type=0,data_id=NULL WHERE id=?", (g["id"],))
            used_counts = {p: 0 for p in PROJECT_INFO}
        assigned = 0
        for g in groups:
            available_types = []
            for ptype in PROJECT_INFO:
                n = c.execute("""SELECT COUNT(*) n FROM project_data pd WHERE pd.project_type=? AND pd.id NOT IN
                    (SELECT data_id FROM groups_tbl WHERE data_id IS NOT NULL AND id<>?)""", (ptype, g["id"])).fetchone()["n"]
                if n > 0:
                    available_types.append(ptype)
            if not available_types:
                continue
            min_count = min(used_counts[p] for p in available_types)
            choices = [p for p in available_types if used_counts[p] == min_count]
            ptype = random.choice(choices)
            c.execute("UPDATE groups_tbl SET project_type=?,data_id=NULL WHERE id=?", (ptype, g["id"]))
            c.execute("UPDATE students SET project_type=? WHERE id IN (SELECT student_id FROM group_students WHERE group_id=?)", (ptype, g["id"]))
            if assign_random_unused_data(c, g["id"], ptype):
                assigned += 1
                used_counts[ptype] += 1
            ensure_progress_rows(c, g["id"])
    return assigned

def ensure_progress_rows(c, group_id: int):
    mids = c.execute("SELECT id FROM milestones ORDER BY position").fetchall()
    c.executemany(
        "INSERT OR IGNORE INTO group_progress(group_id,milestone_id) VALUES (?,?)",
        [(group_id, m["id"]) for m in mids],
    )


def regroup_and_assign():
    """Ghép CHÍNH XÁC theo thứ tự: 1-2, 3-4,... và đọc mã Đề của cặp.
    Hai SV trong một cặp phải cùng mã đề. SV cuối được phép làm cá nhân.
    Hàm kiểm tra toàn bộ trước khi thay phân công cũ.
    """
    with db() as c:
        students = c.execute("SELECT * FROM students ORDER BY stt,id").fetchall()
        if not students:
            raise ValueError("Chưa có danh sách sinh viên.")
        plan=[]
        errors=[]
        for i in range(0, len(students), 2):
            members=list(students[i:i+2])
            ptypes=[int(m["project_type"] or 0) for m in members]
            label=f"SV thứ {i+1}" + (f"-{i+2}" if len(members)==2 else "")
            if any(p not in PROJECT_INFO for p in ptypes):
                errors.append(f"{label}: chưa có mã Đề I-V/1-5.")
                continue
            if len(set(ptypes)) != 1:
                errors.append(f"{label}: hai sinh viên có mã Đề khác nhau ({ptypes[0]} và {ptypes[1]}).")
                continue
            plan.append((members, ptypes[0]))
        if errors:
            raise ValueError("Không thể phân công tự động. " + " | ".join(errors))
        # Kiểm tra đủ dữ liệu cho từng loại đề trước khi ghi
        needs={}
        for _,ptype in plan: needs[ptype]=needs.get(ptype,0)+1
        shortages=[]
        for ptype,n in needs.items():
            have=c.execute("SELECT COUNT(*) n FROM project_data WHERE project_type=?",(ptype,)).fetchone()["n"]
            if have < n: shortages.append(f"DA{ptype}: cần {n} bộ, hiện có {have}")
        if shortages:
            raise ValueError("Không đủ bộ số liệu: " + "; ".join(shortages))
        c.execute("DELETE FROM group_progress")
        c.execute("DELETE FROM group_students")
        c.execute("DELETE FROM groups_tbl")
        for group_no,(members,ptype) in enumerate(plan,1):
            cur=c.execute("INSERT INTO groups_tbl(group_code,project_type) VALUES (?,?)",(f"N{group_no:02d}",ptype))
            gid=cur.lastrowid
            c.executemany("INSERT INTO group_students(group_id,student_id) VALUES (?,?)",[(gid,m["id"]) for m in members])
            assign_next_unused_data(c,gid,ptype)
            ensure_progress_rows(c,gid)
        return len(plan)


def import_students(path: Path):
    rows = parse_students_xlsx(path)
    with db() as c:
        # Import danh sách mới để chờ giảng viên bấm Cập nhật phân công tự động.
        c.execute("DELETE FROM group_progress")
        c.execute("DELETE FROM group_students")
        c.execute("DELETE FROM groups_tbl")
        c.execute("DELETE FROM students")
        c.executemany("INSERT INTO students(stt,student_id,name,project_type) VALUES (?,?,?,?)", rows)
    return len(rows)


def create_manual_group(student1_id:int, student2_id:int|None, project_type:int):
    if project_type not in PROJECT_INFO:
        raise ValueError("Loại đề phải từ I-V (DA1-DA5).")
    ids=[student1_id] + ([student2_id] if student2_id and student2_id!=student1_id else [])
    with db() as c:
        rows=c.execute(f"SELECT * FROM students WHERE id IN ({','.join('?'*len(ids))}) ORDER BY stt,id",ids).fetchall()
        if len(rows)!=len(ids): raise ValueError("Không tìm thấy đủ sinh viên đã chọn.")
        old_group_ids=[]
        for sid in ids:
            old=c.execute("SELECT group_id FROM group_students WHERE student_id=?",(sid,)).fetchone()
            if old:
                old_group_ids.append(old["group_id"]); c.execute("DELETE FROM group_students WHERE student_id=?",(sid,))
        for gid in set(old_group_ids):
            left=c.execute("SELECT COUNT(*) n FROM group_students WHERE group_id=?",(gid,)).fetchone()["n"]
            if left==0: c.execute("DELETE FROM groups_tbl WHERE id=?",(gid,))
        code=next_group_code(c)
        cur=c.execute("INSERT INTO groups_tbl(group_code,project_type) VALUES (?,?)",(code,project_type))
        gid=cur.lastrowid
        c.executemany("INSERT INTO group_students(group_id,student_id) VALUES (?,?)",[(gid,sid) for sid in ids])
        c.execute(f"UPDATE students SET project_type=? WHERE id IN ({','.join('?'*len(ids))})",[project_type,*ids])
        assign_next_unused_data(c,gid,project_type); ensure_progress_rows(c,gid)
        return code

def import_project_data(path: Path):
    rows = parse_project_data(path)
    with db() as c:
        c.execute("UPDATE groups_tbl SET data_id=NULL")
        c.execute("DELETE FROM project_data")
        c.executemany("INSERT INTO project_data(project_type,data_code,p,v,d,t,nam,ngay) VALUES (?,?,?,?,?,?,?,?)", rows)
        # Giữ nguyên phân nhóm thủ công; chỉ cấp lại số liệu theo loại DA hiện tại.
        for ptype in PROJECT_INFO:
            groups = c.execute("SELECT id FROM groups_tbl WHERE project_type=? ORDER BY id", (ptype,)).fetchall()
            for g in groups:
                assign_next_unused_data(c, g["id"], ptype)
                ensure_progress_rows(c, g["id"])
    return len(rows)


def install_template(project_type: int, source_path: Path):
    target = DATA_DIR / f"DA{project_type}_template.docx"
    target.write_bytes(source_path.read_bytes())
    with db() as c:
        row = c.execute("SELECT version FROM templates WHERE project_type=?", (project_type,)).fetchone()
        version = (row["version"] + 1) if row else 1
        c.execute(
            "INSERT INTO templates(project_type,file_path,version,updated_at) VALUES (?,?,?,?) "
            "ON CONFLICT(project_type) DO UPDATE SET file_path=excluded.file_path,version=excluded.version,updated_at=excluded.updated_at",
            (project_type, str(target), version, datetime.now().isoformat(timespec="seconds")),
        )
    return version


def seed_demo_if_empty():
    """Nạp dữ liệu demo chỉ khi bộ demo thực sự có trong bản deploy.

    Trên GitHub/Render, nếu người dùng vô tình chưa đẩy thư mục demo thì ứng dụng
    vẫn phải khởi động để giảng viên có thể import dữ liệu thật từ giao diện.
    """
    with db() as c:
        if c.execute("SELECT COUNT(*) n FROM students").fetchone()["n"]:
            return
    required = [DEMO_DIR / "dataset.xlsx", DEMO_DIR / "student_list.xlsx"]
    required += [DEMO_DIR / "templates" / f"DA{ptype}_PBL1.docx" for ptype in PROJECT_INFO]
    if not all(fp.exists() for fp in required):
        return
    for ptype in PROJECT_INFO:
        install_template(ptype, DEMO_DIR / "templates" / f"DA{ptype}_PBL1.docx")
    import_project_data(DEMO_DIR / "dataset.xlsx")
    import_students(DEMO_DIR / "student_list.xlsx")
    try:
        regroup_and_assign()
    except Exception:
        pass


def get_settings():
    with db() as c:
        rows = c.execute("SELECT key,value FROM app_settings").fetchall()
    return {r["key"]: r["value"] for r in rows}


def update_settings(instructor_name: str, start_week: str, end_week: str, academic_year_start: str, academic_year_end: str, attendance_weight: str = "10"):
    values = {
        "instructor_name": instructor_name.strip(),
        "start_week": start_week.strip(),
        "end_week": end_week.strip(),
        "academic_year_start": academic_year_start.strip(),
        "academic_year_end": academic_year_end.strip(),
        "attendance_weight": attendance_weight.strip() or "10",
    }
    for key in ("start_week", "end_week", "academic_year_start", "academic_year_end", "attendance_weight"):
        if values[key] and not values[key].isdigit():
            raise ValueError(f"{key} phải là số.")
    if values["start_week"] and values["end_week"] and int(values["start_week"]) > int(values["end_week"]):
        raise ValueError("Tuần bắt đầu không được lớn hơn tuần kết thúc.")
    with db() as c:
        for key, value in values.items():
            c.execute("INSERT INTO app_settings(key,value) VALUES (?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, value))


def load_image_url(project_type: int):
    return f"/static/project_images/DA{project_type}_load.png"


def system_image_url(project_type: int):
    return f"/static/project_images/DA{project_type}_system.png"


def format_num(v):
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


def set_para_text(paragraph, text):
    if paragraph.runs:
        for r in paragraph.runs:
            r.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def generate_docx(group_id: int):
    with db() as c:
        g = c.execute(
            """SELECT g.*,pd.data_code,pd.p,pd.v,pd.d,pd.t,pd.nam,pd.ngay,t.file_path,t.version
               FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id
               LEFT JOIN templates t ON t.project_type=g.project_type WHERE g.id=?""",
            (group_id,),
        ).fetchone()
        if not g:
            raise ValueError("Không tìm thấy nhóm")
        members = c.execute("SELECT s.* FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (group_id,)).fetchall()
    if not g["file_path"]:
        raise ValueError("Chưa có template Word cho loại đề này")
    if not g["data_code"]:
        raise ValueError("Nhóm chưa được cấp dòng số liệu")

    doc = Document(g["file_path"])
    names = " - ".join(m["name"] for m in members)
    if len(doc.tables) > 1:
        info = doc.tables[1]
        set_para_text(info.cell(0, 1).paragraphs[0], names)
        set_para_text(info.cell(1, 1).paragraphs[0], f"{g['group_code']} / {g['data_code']}")

    p, v, d, tval, nam, ngay = [format_num(g[k]) for k in ("p", "v", "d", "t", "nam", "ngay")]
    if g["project_type"] in (1, 2, 3, 4) and len(doc.tables) > 2:
        cell = doc.tables[2].cell(1, 0)
        paras = cell.paragraphs
        labels = {
            1: ("Lực kéo băng tải", "Vận tốc băng tải"),
            2: ("Lực kéo dây cáp", "Vận tốc kéo cáp"),
            3: ("Lực kéo băng tải", "Vận tốc băng tải"),
            4: ("Lực kéo cáp", "Vận tốc kéo cáp"),
        }[g["project_type"]]
        lines = [f"1. {labels[0]} :    P = {p}", f"2. {labels[1]} :    V = {v}", f"3. Đường kính tang :    D = {d}"]
        for i in range(min(3, len(paras))):
            set_para_text(paras[i], lines[i])
        for para in paras:
            txt = para.text
            if "Thời gian phục vụ" in txt:
                set_para_text(para, f"5. Thời gian phục vụ: {tval} năm")
            elif "Một năm làm việc" in txt:
                set_para_text(para, f"Một năm làm việc {nam} ngày, một ngày làm việc {ngay} giờ")
    elif g["project_type"] == 5:
        for para in doc.paragraphs:
            txt = para.text
            if txt.startswith("1. Lực kéo băng tải"):
                set_para_text(para, f"1. Lực kéo băng tải :\t\tP = {p}")
            elif txt.startswith("2. Vận tốc băng tải"):
                set_para_text(para, f"2. Vận tốc băng tải :\t\tV = {v}")
            elif txt.startswith("3. Đường kính tang"):
                set_para_text(para, f"3. Đường kính tang :\t\tD = {d}")
            elif txt.startswith("5. Thời gian phục vụ"):
                set_para_text(para, f"5. Thời gian phục vụ: {tval} năm")
            elif txt.startswith("Một năm làm việc"):
                set_para_text(para, f"Một năm làm việc {nam} ngày, một ngày làm việc {ngay} giờ.")

    out = GENERATED_DIR / f"{g['group_code']}_DA{g['project_type']}_{g['data_code'].replace('.', '-')}.docx"
    doc.save(out)
    return out


ALLOWED_SUBMISSION_EXTS = {".docx", ".xlsx", ".xls", ".dwg", ".dxf"}

def safe_filename(name: str):
    import re
    base = Path(name).name
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", base)
    return stem[:160] or "file"

def group_members_text(c, group_id: int):
    rows = c.execute("SELECT s.name FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (group_id,)).fetchall()
    return " - ".join(r["name"] for r in rows)


def render(request: Request, name: str, **kwargs):
    return templates.TemplateResponse(request=request, name=name, context={"PROJECT_INFO": PROJECT_INFO, "format_num": format_num, "settings": get_settings(), "staff": current_staff(request), "load_image_url": load_image_url, "system_image_url": system_image_url, **kwargs})


def redirect_with(path: str, *, msg: str | None = None, err: str | None = None):
    from urllib.parse import urlencode
    q = urlencode({k: v for k, v in {"msg": msg, "err": err}.items() if v})
    return RedirectResponse(path + ("?" + q if q else ""), status_code=303)


@app.get("/admin-login", response_class=HTMLResponse, name="admin_login")
def admin_login_get(request: Request):
    return render(request, "admin_login.html")

@app.post("/admin-login", name="admin_login_post")
def admin_login_post(username: str = Form("admin"), password: str = Form(...)):
    username = username.strip() or "admin"
    if username.lower() == "admin":
        if not secrets.compare_digest(password, ADMIN_PASSWORD):
            return redirect_with("/admin-login", err="Tài khoản hoặc mật khẩu không đúng.")
        cookie = f"admin:0:{_sign('admin:0')}"
    else:
        with db() as c:
            row = c.execute("SELECT * FROM staff_users WHERE username=? AND active=1", (username,)).fetchone()
        if not row or not verify_password(password, row["password_hash"]):
            return redirect_with("/admin-login", err="Tài khoản hoặc mật khẩu không đúng.")
        cookie = f"reviewer:{row['id']}:{_sign(f"reviewer:{row['id']}")}"
    response = RedirectResponse(url="/", status_code=303)
    response.set_cookie("pbl_staff", cookie, httponly=True, samesite="lax", max_age=60*60*12)
    return response

@app.get("/admin-logout", name="admin_logout")
def admin_logout():
    response = RedirectResponse(url="/admin-login", status_code=303)
    response.delete_cookie("pbl_staff")
    return response


@app.get("/healthz", include_in_schema=False)
def healthz():
    return {"status": "ok"}

@app.on_event("startup")
def startup():
    init_db()
    seed_demo_if_empty()


@app.get("/", response_class=HTMLResponse, name="dashboard")
def dashboard(request: Request):
    with db() as c:
        stats = {
            "students": c.execute("SELECT COUNT(*) n FROM students").fetchone()["n"],
            "groups": c.execute("SELECT COUNT(*) n FROM groups_tbl").fetchone()["n"],
            "assigned": c.execute("SELECT COUNT(*) n FROM groups_tbl WHERE data_id IS NOT NULL").fetchone()["n"],
            "templates": c.execute("SELECT COUNT(*) n FROM templates").fetchone()["n"],
        }
        groups = c.execute(
            """SELECT g.*,pd.data_code,
                      (SELECT GROUP_CONCAT(s.name, ' • ') FROM group_students gs JOIN students s ON s.id=gs.student_id WHERE gs.group_id=g.id) members,
                      COALESCE((SELECT ROUND(100.0*SUM(CASE WHEN gp.status='Hoàn thành' THEN 1 ELSE 0 END)/NULLIF(COUNT(*),0)) FROM group_progress gp WHERE gp.group_id=g.id),0) progress
               FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id ORDER BY g.id"""
        ).fetchall()
        unassigned = c.execute("SELECT s.* FROM students s WHERE NOT EXISTS (SELECT 1 FROM group_students gs WHERE gs.student_id=s.id) ORDER BY s.stt,s.id").fetchall()
        all_students = c.execute("SELECT s.*, (SELECT g.group_code FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=s.id LIMIT 1) group_code FROM students s ORDER BY s.stt,s.id").fetchall()
    return render(request, "dashboard.html", stats=stats, groups=groups, unassigned=unassigned, all_students=all_students)


@app.get("/admin/import", response_class=HTMLResponse, name="admin_import")
def admin_import_get(request: Request):
    with db() as c:
        trows = c.execute("SELECT * FROM templates ORDER BY project_type").fetchall()
        template_map = {r["project_type"]: r for r in trows}
        counts = c.execute("SELECT project_type,COUNT(*) n FROM project_data GROUP BY project_type ORDER BY project_type").fetchall()
    return render(request, "import.html", template_map=template_map, counts=counts)


@app.post("/admin/import", name="admin_import_post")
async def admin_import_post(kind: str = Form(...), file: UploadFile = File(...), project_type: int | None = Form(None)):
    suffix = Path(file.filename or "").suffix.lower()
    tmp = DATA_DIR / f"upload_{datetime.now().strftime('%Y%m%d%H%M%S%f')}{suffix}"
    try:
        tmp.write_bytes(await file.read())
        if kind == "students":
            if suffix != ".xlsx":
                raise ValueError("Danh sách sinh viên V0.1 yêu cầu file .xlsx")
            n = import_students(tmp)
            return redirect_with("/admin/import", msg=f"Đã import {n} sinh viên. Vào trang Phân công để cập nhật tự động theo mã Đề hoặc phân công thủ công.")
        if kind == "dataset":
            n = import_project_data(tmp)
            return redirect_with("/admin/import", msg=f"Đã import {n} dòng số liệu cho DA1-DA5.")
        if kind == "template":
            if project_type not in PROJECT_INFO or suffix != ".docx":
                raise ValueError("Template phải là .docx và DA từ 1 đến 5.")
            version = install_template(int(project_type), tmp)
            return redirect_with("/admin/import", msg=f"Đã cập nhật template DA{project_type}, phiên bản {version}.")
        raise ValueError("Loại import không hợp lệ")
    except Exception as exc:
        return redirect_with("/admin/import", err=str(exc))
    finally:
        if tmp.exists():
            tmp.unlink()


@app.post("/admin/settings", name="admin_settings_post")
def admin_settings_post(
    instructor_name: str = Form(""),
    start_week: str = Form(""),
    end_week: str = Form(""),
    academic_year_start: str = Form(""),
    academic_year_end: str = Form(""),
    attendance_weight: str = Form("10"),
):
    try:
        update_settings(instructor_name, start_week, end_week, academic_year_start, academic_year_end, attendance_weight)
    except Exception as exc:
        return redirect_with("/admin/import", err=str(exc))
    return redirect_with("/admin/import", msg="Đã lưu thông tin học phần và giảng viên hướng dẫn.")


@app.get("/admin/groups", response_class=HTMLResponse, name="admin_groups")
def admin_groups(request: Request):
    with db() as c:
        groups = c.execute("""SELECT g.*,pd.data_code,
            (SELECT GROUP_CONCAT(s.name, ' • ') FROM group_students gs JOIN students s ON s.id=gs.student_id WHERE gs.group_id=g.id) members,
            (SELECT COUNT(*) FROM group_students gs WHERE gs.group_id=g.id) member_count
            FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id ORDER BY g.group_code""").fetchall()
        students = c.execute("""SELECT s.*, (SELECT g.id FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=s.id LIMIT 1) group_id,
            (SELECT g.group_code FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=s.id LIMIT 1) group_code
            FROM students s ORDER BY s.stt,s.id""").fetchall()
        data_rows = c.execute("SELECT id,project_type,data_code FROM project_data ORDER BY project_type,id").fetchall()
    return render(request, "admin_groups.html", groups=groups, students=students, data_rows=data_rows)

@app.post("/admin/groups/auto", name="admin_groups_auto")
def admin_groups_auto():
    try:
        n=regroup_and_assign()
        return redirect_with("/admin/groups", msg=f"Đã cập nhật phân công tự động: {n} nhóm/cá nhân theo thứ tự danh sách.")
    except Exception as exc:
        return redirect_with("/admin/groups", err=str(exc))


@app.post("/admin/groups/manual", name="admin_groups_manual")
def admin_groups_manual(student1_id:int=Form(...), student2_id:int=Form(0), project_type:int=Form(...)):
    try:
        code=create_manual_group(student1_id, student2_id or None, project_type)
        return redirect_with("/admin/groups", msg=f"Đã tạo/cập nhật phân công thủ công {code}.")
    except Exception as exc:
        return redirect_with("/admin/groups", err=str(exc))



@app.post("/admin/groups/move", name="admin_group_move")
def admin_group_move(student_db_id: int = Form(...), target_group_id: int = Form(0), new_project_type: int = Form(0)):
    try:
        with db() as c:
            student = c.execute("SELECT * FROM students WHERE id=?", (student_db_id,)).fetchone()
            if not student:
                raise ValueError("Không tìm thấy sinh viên.")
            old = c.execute("SELECT g.id FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=?", (student_db_id,)).fetchone()
            if target_group_id:
                target = c.execute("SELECT * FROM groups_tbl WHERE id=?", (target_group_id,)).fetchone()
                if not target:
                    raise ValueError("Không tìm thấy nhóm đích.")
                count = c.execute("SELECT COUNT(*) n FROM group_students WHERE group_id=?", (target_group_id,)).fetchone()["n"]
                if count >= 2 and (not old or old["id"] != target_group_id):
                    raise ValueError("Nhóm đích đã có 2 sinh viên.")
                if old and old["id"] == target_group_id:
                    return redirect_with("/admin/groups", msg="Sinh viên đã ở nhóm này.")
                if old:
                    c.execute("DELETE FROM group_students WHERE group_id=? AND student_id=?", (old["id"], student_db_id))
                c.execute("INSERT OR IGNORE INTO group_students(group_id,student_id) VALUES (?,?)", (target_group_id, student_db_id))
                c.execute("UPDATE students SET project_type=? WHERE id=?", (target["project_type"], student_db_id))
            else:
                if new_project_type not in (0,1,2,3,4,5):
                    raise ValueError("Loại đề không hợp lệ.")
                if old:
                    c.execute("DELETE FROM group_students WHERE group_id=? AND student_id=?", (old["id"], student_db_id))
                code = next_group_code(c)
                cur = c.execute("INSERT INTO groups_tbl(group_code,project_type) VALUES (?,?)", (code, new_project_type))
                target_group_id = cur.lastrowid
                c.execute("INSERT INTO group_students(group_id,student_id) VALUES (?,?)", (target_group_id, student_db_id))
                c.execute("UPDATE students SET project_type=? WHERE id=?", (new_project_type, student_db_id))
                if new_project_type in PROJECT_INFO:
                    assign_next_unused_data(c, target_group_id, new_project_type)
                ensure_progress_rows(c, target_group_id)
            if old:
                nleft = c.execute("SELECT COUNT(*) n FROM group_students WHERE group_id=?", (old["id"],)).fetchone()["n"]
                if nleft == 0:
                    c.execute("DELETE FROM groups_tbl WHERE id=?", (old["id"],))
            ensure_progress_rows(c, target_group_id)
        return redirect_with("/admin/groups", msg="Đã cập nhật phân công sinh viên/nhóm.")
    except Exception as exc:
        return redirect_with("/admin/groups", err=str(exc))

@app.post("/admin/groups/{group_id}/project", name="admin_group_project")
def admin_group_project(group_id: int, project_type: int = Form(...)):
    try:
        if project_type not in PROJECT_INFO:
            raise ValueError("Loại đề phải từ DA1 đến DA5.")
        with db() as c:
            g = c.execute("SELECT * FROM groups_tbl WHERE id=?", (group_id,)).fetchone()
            if not g:
                raise ValueError("Không tìm thấy nhóm.")
            c.execute("UPDATE groups_tbl SET project_type=?, data_id=NULL WHERE id=?", (project_type, group_id))
            c.execute("UPDATE students SET project_type=? WHERE id IN (SELECT student_id FROM group_students WHERE group_id=?)", (project_type, group_id))
            assign_next_unused_data(c, group_id, project_type)
        return redirect_with("/admin/groups", msg="Đã đổi loại đề và cấp lại bộ số liệu phù hợp.")
    except Exception as exc:
        return redirect_with("/admin/groups", err=str(exc))


@app.get("/admin/plan", response_class=HTMLResponse, name="admin_plan")
def admin_plan(request: Request):
    with db() as c:
        rows=c.execute("SELECT * FROM milestones ORDER BY position,id").fetchall()
    return render(request,"admin_plan.html",milestones=rows)


@app.post("/admin/plan/{milestone_id}/save", name="admin_plan_save")
def admin_plan_save(milestone_id:int, title:str=Form(...), description:str=Form(""), start_at:str=Form(""), due_at:str=Form("")):
    try:
        with db() as c:
            if not c.execute("SELECT id FROM milestones WHERE id=?",(milestone_id,)).fetchone(): raise ValueError("Không tìm thấy nội dung thực hiện.")
            c.execute("UPDATE milestones SET title=?,description=?,start_at=?,due_at=? WHERE id=?",(title.strip(),description.strip(),start_at.strip(),due_at.strip(),milestone_id))
        return redirect_with("/admin/plan",msg="Đã lưu nội dung và mốc thời gian.")
    except Exception as exc:
        return redirect_with("/admin/plan",err=str(exc))


@app.post("/admin/plan/{milestone_id}/activate", name="admin_plan_activate")
def admin_plan_activate(milestone_id:int):
    try:
        with db() as c:
            if not c.execute("SELECT id FROM milestones WHERE id=?",(milestone_id,)).fetchone(): raise ValueError("Không tìm thấy nội dung thực hiện.")
            c.execute("UPDATE milestones SET is_active=CASE WHEN id=? THEN 1 ELSE 0 END",(milestone_id,))
            group_ids=c.execute("SELECT id FROM groups_tbl").fetchall()
            for g in group_ids:
                ensure_progress_rows(c,g["id"])
                c.execute("UPDATE group_progress SET status='Đang thực hiện' WHERE group_id=? AND milestone_id=? AND status<>'Hoàn thành'",(g["id"],milestone_id))
        return redirect_with("/admin/plan",msg="Đã kích hoạt nội dung cho TẤT CẢ nhóm/cá nhân.")
    except Exception as exc:
        return redirect_with("/admin/plan",err=str(exc))


@app.post("/admin/plan/{milestone_id}/close", name="admin_plan_close")
def admin_plan_close(milestone_id:int):
    with db() as c:
        c.execute("UPDATE milestones SET is_active=0 WHERE id=?",(milestone_id,))
    return redirect_with("/admin/plan",msg="Đã đóng nội dung thực hiện.")



@app.get("/group/{group_id}", response_class=HTMLResponse, name="group_detail")
def group_detail(request: Request, group_id: int):
    with db() as c:
        g = c.execute("SELECT g.*,pd.data_code,pd.p,pd.v,pd.d,pd.t,pd.nam,pd.ngay FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id WHERE g.id=?", (group_id,)).fetchone()
        if not g:
            return PlainTextResponse("Không tìm thấy nhóm", status_code=404)
        members = c.execute("SELECT s.* FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (group_id,)).fetchall()
        progress = c.execute("SELECT m.*,gp.status,gp.score,gp.feedback FROM milestones m LEFT JOIN group_progress gp ON gp.milestone_id=m.id AND gp.group_id=? ORDER BY m.position", (group_id,)).fetchall()
        submissions = c.execute("""SELECT sub.*,s.name student_name,m.title milestone_title FROM submissions sub
            JOIN students s ON s.id=sub.student_id LEFT JOIN milestones m ON m.id=sub.milestone_id
            WHERE sub.group_id=? ORDER BY sub.uploaded_at DESC""", (group_id,)).fetchall()
        evaluations = c.execute("""SELECT e.*,m.title milestone_title FROM evaluation_log e JOIN milestones m ON m.id=e.milestone_id
            WHERE e.group_id=? ORDER BY e.id DESC LIMIT 50""", (group_id,)).fetchall()
    attendance = attendance_summary_for_students([m["id"] for m in members])
    return render(request, "group.html", g=g, members=members, progress=progress, submissions=submissions, evaluations=evaluations, attendance=attendance, info=PROJECT_INFO.get(g["project_type"], {"title":"Chưa phân đề","task":"Nhóm chưa được phân loại dự án."}))


@app.post("/group/{group_id}/progress/{milestone_id}", name="update_progress")
def update_progress(request: Request, group_id: int, milestone_id: int, status: str = Form(...), score: str = Form(""), feedback: str = Form("")):
    try:
        score_value = float(score) if score.strip() else None
    except ValueError:
        return redirect_with(f"/group/{group_id}", err="Điểm phải là số.")
    evaluator = (current_staff(request) or {}).get("name", "Giảng viên")
    with db() as c:
        c.execute("INSERT INTO group_progress(group_id,milestone_id,status,score,feedback) VALUES (?,?,?,?,?) ON CONFLICT(group_id,milestone_id) DO UPDATE SET status=excluded.status,score=excluded.score,feedback=excluded.feedback", (group_id, milestone_id, status, score_value, feedback.strip()))
        c.execute("INSERT INTO evaluation_log(group_id,milestone_id,evaluator,score,feedback,created_at) VALUES (?,?,?,?,?,?)",
                  (group_id, milestone_id, evaluator, score_value, feedback.strip(), datetime.now().isoformat(timespec="seconds")))
    return redirect_with(f"/group/{group_id}", msg=f"Đã lưu đánh giá của {evaluator}.")


@app.get("/student", response_class=HTMLResponse, name="student_lookup")
def student_lookup_get(request: Request):
    return render(request, "student_lookup.html")


@app.post("/student", name="student_lookup_post")
def student_lookup_post(student_id: str = Form(...)):
    sid = student_id.strip()
    with db() as c:
        row = c.execute("SELECT s.student_id FROM students s WHERE s.student_id=?", (sid,)).fetchone()
    if not row:
        return redirect_with("/student", err="Không tìm thấy số thẻ sinh viên.")
    return RedirectResponse(f"/student/{sid}", status_code=303)


@app.get("/student/{student_id}", response_class=HTMLResponse, name="student_portal")
def student_portal(request: Request, student_id: str):
    with db() as c:
        s = c.execute("SELECT * FROM students WHERE student_id=?", (student_id,)).fetchone()
        if not s:
            return PlainTextResponse("Không tìm thấy sinh viên", status_code=404)
        g = c.execute("SELECT g.*,pd.data_code,pd.p,pd.v,pd.d,pd.t,pd.nam,pd.ngay FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id LEFT JOIN project_data pd ON pd.id=g.data_id WHERE gs.student_id=?", (s["id"],)).fetchone()
        if not g or int(g["project_type"] or 0) not in PROJECT_INFO or not g["data_code"]:
            return render(request, "student_unassigned.html", s=s)
        members = c.execute("SELECT s.* FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (g["id"],)).fetchall()
        progress = c.execute("SELECT m.*,gp.status,gp.score,gp.feedback FROM milestones m LEFT JOIN group_progress gp ON gp.milestone_id=m.id AND gp.group_id=? ORDER BY m.position", (g["id"],)).fetchall()
        submissions = c.execute("""SELECT sub.*,s.name student_name,m.title milestone_title FROM submissions sub
            JOIN students s ON s.id=sub.student_id LEFT JOIN milestones m ON m.id=sub.milestone_id
            WHERE sub.group_id=? ORDER BY sub.uploaded_at DESC""", (g["id"],)).fetchall()
    done = sum(1 for x in progress if x["status"] == "Hoàn thành")
    percent = round(100 * done / len(progress)) if progress else 0
    with db() as c:
        materials = c.execute("SELECT * FROM materials ORDER BY created_at DESC").fetchall()
    attendance = attendance_summary_for_students([s["id"]]).get(s["id"], {"present":0,"total":0,"score":0})
    info = PROJECT_INFO.get(g["project_type"], {"title":"Chưa phân đề","task":"Nhóm chưa được phân đề."})
    return render(request, "student_portal.html", s=s, g=g, members=members, progress=progress, submissions=submissions, materials=materials, attendance=attendance, percent=percent, info=info)


@app.post("/student/{student_id}/submit", name="student_submit")
async def student_submit(student_id: str, milestone_id: int = Form(...), note: str = Form(""), files: list[UploadFile] = File(...)):
    try:
        with db() as c:
            srow = c.execute("SELECT * FROM students WHERE student_id=?", (student_id,)).fetchone()
            if not srow:
                raise ValueError("Không tìm thấy sinh viên.")
            grow = c.execute("SELECT g.* FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=?", (srow["id"],)).fetchone()
            if not grow:
                raise ValueError("Sinh viên chưa được phân nhóm.")
            valid_mid = c.execute("SELECT id,is_active,due_at FROM milestones WHERE id=?", (milestone_id,)).fetchone()
            if not valid_mid:
                raise ValueError("Mốc tiến độ không hợp lệ.")
            if not valid_mid["is_active"]:
                raise ValueError("Nội dung này chưa được giảng viên kích hoạt hoặc đã đóng.")
            saved = 0
            for upload in files:
                original = upload.filename or ""
                ext = Path(original).suffix.lower()
                if ext not in ALLOWED_SUBMISSION_EXTS:
                    raise ValueError(f"Không hỗ trợ file {original}. Chỉ nhận DOCX, XLS/XLSX, DWG/DXF.")
                content = await upload.read()
                if len(content) > 50 * 1024 * 1024:
                    raise ValueError(f"File {original} vượt quá 50 MB.")
                stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                stored = f"G{grow['id']}_S{srow['id']}_{stamp}_{safe_filename(original)}"
                (SUBMISSIONS_DIR / stored).write_bytes(content)
                c.execute("""INSERT INTO submissions(group_id,student_id,milestone_id,note,original_name,stored_name,file_type,uploaded_at)
                    VALUES (?,?,?,?,?,?,?,?)""", (grow["id"], srow["id"], milestone_id, note.strip(), original, stored, ext, datetime.now().isoformat(timespec="seconds")))
                saved += 1
            c.execute("""INSERT INTO group_progress(group_id,milestone_id,status) VALUES (?,?,?)
                ON CONFLICT(group_id,milestone_id) DO UPDATE SET status=CASE WHEN group_progress.status='Hoàn thành' THEN group_progress.status ELSE excluded.status END""",
                (grow["id"], milestone_id, "Đang thực hiện"))
        return redirect_with(f"/student/{student_id}", msg=f"Đã gửi báo cáo tiến độ và tải lên {saved} file.")
    except Exception as exc:
        return redirect_with(f"/student/{student_id}", err=str(exc))

@app.get("/submission/{submission_id}/download", name="download_submission")
def download_submission(submission_id: int):
    with db() as c:
        row = c.execute("SELECT * FROM submissions WHERE id=?", (submission_id,)).fetchone()
    if not row:
        return PlainTextResponse("Không tìm thấy file", status_code=404)
    path = SUBMISSIONS_DIR / row["stored_name"]
    if not path.exists():
        return PlainTextResponse("File không còn trên máy chủ", status_code=404)
    return FileResponse(path, filename=row["original_name"])



def attendance_summary_for_students(student_db_ids):
    ids = [int(x) for x in student_db_ids if x is not None]
    if not ids:
        return {}
    ph = ",".join("?" for _ in ids)
    with db() as c:
        total = c.execute("SELECT COUNT(*) n FROM attendance_sessions WHERE is_open=0").fetchone()["n"]
        rows = c.execute(f"""SELECT student_id,COUNT(*) n FROM attendance_records ar
            JOIN attendance_sessions a ON a.id=ar.session_id
            WHERE a.is_open=0 AND ar.status='Có mặt' AND student_id IN ({ph}) GROUP BY student_id""", ids).fetchall()
    present = {r["student_id"]: r["n"] for r in rows}
    return {sid: {"present": present.get(sid,0), "total": total, "score": round(10*present.get(sid,0)/total,2) if total else 0} for sid in ids}

@app.post("/admin/groups/random", name="admin_groups_random")
def admin_groups_random(scope: str = Form("unassigned")):
    try:
        n = random_assign_groups(only_unassigned=(scope != "all"))
        return redirect_with("/admin/groups", msg=f"Đã phân ngẫu nhiên đề và bộ số liệu cho {n} nhóm.")
    except Exception as exc:
        return redirect_with("/admin/groups", err=str(exc))

@app.get("/admin/reviewers", response_class=HTMLResponse, name="admin_reviewers")
def admin_reviewers(request: Request):
    with db() as c:
        rows = c.execute("SELECT id,username,display_name,role,active,created_at FROM staff_users ORDER BY display_name").fetchall()
    return render(request, "admin_reviewers.html", reviewers=rows)

@app.post("/admin/reviewers", name="admin_reviewer_add")
def admin_reviewer_add(username: str = Form(...), display_name: str = Form(...), password: str = Form(...)):
    try:
        username = username.strip()
        if len(username) < 3 or len(password) < 6:
            raise ValueError("Tên đăng nhập tối thiểu 3 ký tự; mật khẩu tối thiểu 6 ký tự.")
        with db() as c:
            c.execute("INSERT INTO staff_users(username,display_name,password_hash,role,active,created_at) VALUES (?,?,?,?,1,?)",
                      (username, display_name.strip(), hash_password(password), "reviewer", datetime.now().isoformat(timespec="seconds")))
        return redirect_with("/admin/reviewers", msg="Đã tạo tài khoản đồng nghiệp.")
    except Exception as exc:
        return redirect_with("/admin/reviewers", err=str(exc))

@app.post("/admin/reviewers/{user_id}/toggle", name="admin_reviewer_toggle")
def admin_reviewer_toggle(user_id: int):
    with db() as c:
        c.execute("UPDATE staff_users SET active=CASE active WHEN 1 THEN 0 ELSE 1 END WHERE id=?", (user_id,))
    return redirect_with("/admin/reviewers", msg="Đã cập nhật trạng thái tài khoản.")

ALLOWED_MATERIAL_EXTS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".zip", ".rar", ".dwg", ".dxf", ".txt"}

@app.get("/materials", response_class=HTMLResponse, name="materials_page")
def materials_page(request: Request):
    with db() as c:
        rows = c.execute("SELECT * FROM materials ORDER BY created_at DESC").fetchall()
    return render(request, "materials.html", materials=rows)

@app.get("/admin/materials", response_class=HTMLResponse, name="admin_materials")
def admin_materials(request: Request):
    with db() as c:
        rows = c.execute("SELECT * FROM materials ORDER BY created_at DESC").fetchall()
    return render(request, "admin_materials.html", materials=rows)

@app.post("/admin/materials/link", name="admin_material_link")
def admin_material_link(title: str = Form(...), description: str = Form(""), url: str = Form(...)):
    if not url.strip().lower().startswith(("http://", "https://")):
        return redirect_with("/admin/materials", err="Link phải bắt đầu bằng http:// hoặc https://")
    with db() as c:
        c.execute("INSERT INTO materials(title,description,kind,url,created_at) VALUES (?,?,?,?,?)",
                  (title.strip(), description.strip(), "link", url.strip(), datetime.now().isoformat(timespec="seconds")))
    return redirect_with("/admin/materials", msg="Đã thêm liên kết tài liệu.")

@app.post("/admin/materials/file", name="admin_material_file")
async def admin_material_file(title: str = Form(...), description: str = Form(""), file: UploadFile = File(...)):
    try:
        original = file.filename or ""
        ext = Path(original).suffix.lower()
        if ext not in ALLOWED_MATERIAL_EXTS:
            raise ValueError("Định dạng tài liệu chưa được hỗ trợ.")
        content = await file.read()
        if len(content) > 80*1024*1024:
            raise ValueError("Tài liệu vượt quá 80 MB.")
        stored = f"{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}_{safe_filename(original)}"
        (MATERIALS_DIR/stored).write_bytes(content)
        with db() as c:
            c.execute("INSERT INTO materials(title,description,kind,original_name,stored_name,created_at) VALUES (?,?,?,?,?,?)",
                      (title.strip(), description.strip(), "file", original, stored, datetime.now().isoformat(timespec="seconds")))
        return redirect_with("/admin/materials", msg="Đã tải tài liệu lên.")
    except Exception as exc:
        return redirect_with("/admin/materials", err=str(exc))

@app.get("/material/{material_id}/download", name="material_download")
def material_download(material_id: int):
    with db() as c:
        row = c.execute("SELECT * FROM materials WHERE id=?", (material_id,)).fetchone()
    if not row:
        return PlainTextResponse("Không tìm thấy tài liệu", status_code=404)
    if row["kind"] == "link":
        return RedirectResponse(row["url"], status_code=302)
    path = MATERIALS_DIR / row["stored_name"]
    if not path.exists():
        return PlainTextResponse("File tài liệu không còn trên máy chủ", status_code=404)
    return FileResponse(path, filename=row["original_name"])

@app.post("/admin/materials/{material_id}/delete", name="admin_material_delete")
def admin_material_delete(material_id: int):
    with db() as c:
        row = c.execute("SELECT * FROM materials WHERE id=?", (material_id,)).fetchone()
        if row:
            c.execute("DELETE FROM materials WHERE id=?", (material_id,))
    if row and row["stored_name"]:
        path = MATERIALS_DIR / row["stored_name"]
        if path.exists(): path.unlink()
    return redirect_with("/admin/materials", msg="Đã xóa tài liệu.")

@app.get("/admin/backup", name="admin_backup")
def admin_backup():
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = BACKUPS_DIR / f"PBL_backup_{stamp}.zip"
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        if DB_PATH.exists(): z.write(DB_PATH, "data/pbl.db")
        for folder_name, folder in (("submissions", SUBMISSIONS_DIR), ("materials", MATERIALS_DIR)):
            for fp in folder.rglob("*"):
                if fp.is_file(): z.write(fp, f"data/{folder_name}/{fp.name}")
        for ptype in PROJECT_INFO:
            fp = DATA_DIR / f"DA{ptype}_template.docx"
            if fp.exists(): z.write(fp, f"data/{fp.name}")
    return FileResponse(out, media_type="application/zip", filename=out.name)

@app.post("/admin/restore", name="admin_restore")
async def admin_restore(file: UploadFile = File(...)):
    if Path(file.filename or "").suffix.lower() != ".zip":
        return redirect_with("/admin/import", err="Bản sao lưu phải là file .zip")
    tmpzip = BACKUPS_DIR / f"restore_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.zip"
    tmpdir = BACKUPS_DIR / (tmpzip.stem + "_dir")
    try:
        tmpzip.write_bytes(await file.read())
        tmpdir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(tmpzip) as z:
            for member in z.infolist():
                dest = (tmpdir / member.filename).resolve()
                if not str(dest).startswith(str(tmpdir.resolve())):
                    raise ValueError("File backup không hợp lệ.")
            z.extractall(tmpdir)
        srcdb = tmpdir / "data" / "pbl.db"
        if not srcdb.exists():
            raise ValueError("Backup không chứa data/pbl.db")
        shutil.copy2(srcdb, DB_PATH)
        for name, target in (("submissions", SUBMISSIONS_DIR), ("materials", MATERIALS_DIR)):
            src = tmpdir / "data" / name
            if src.exists():
                target.mkdir(exist_ok=True)
                for fp in src.iterdir():
                    if fp.is_file(): shutil.copy2(fp, target/fp.name)
        for ptype in PROJECT_INFO:
            src = tmpdir / "data" / f"DA{ptype}_template.docx"
            if src.exists(): shutil.copy2(src, DATA_DIR/src.name)
        init_db()
        return redirect_with("/admin/import", msg="Đã khôi phục dữ liệu. Phân nhóm, đề, điểm và tài liệu đã được nạp lại.")
    except Exception as exc:
        return redirect_with("/admin/import", err=str(exc))
    finally:
        if tmpzip.exists(): tmpzip.unlink()
        if tmpdir.exists(): shutil.rmtree(tmpdir, ignore_errors=True)

@app.get("/admin/attendance", response_class=HTMLResponse, name="admin_attendance")
def admin_attendance(request: Request):
    with db() as c:
        sessions = c.execute("""SELECT a.*, (SELECT COUNT(*) FROM attendance_records ar WHERE ar.session_id=a.id AND ar.status='Có mặt') present_count FROM attendance_sessions a ORDER BY a.id DESC""").fetchall()
    return render(request, "attendance.html", sessions=sessions)

@app.post("/admin/attendance/create", name="attendance_create")
def attendance_create(title: str = Form(...), duration_minutes: int = Form(10)):
    duration_minutes = max(1, min(int(duration_minutes), 180))
    now = datetime.now()
    token = secrets.token_urlsafe(18)
    with db() as c:
        c.execute("INSERT INTO attendance_sessions(title,session_date,token,expires_at,is_open,created_at) VALUES (?,?,?,?,1,?)",
                  (title.strip(), now.strftime("%d/%m/%Y %H:%M"), token, (now+timedelta(minutes=duration_minutes)).isoformat(timespec="seconds"), now.isoformat(timespec="seconds")))
    return redirect_with("/admin/attendance", msg="Đã mở phiên điểm danh QR.")

@app.post("/admin/attendance/{session_id}/close", name="attendance_close")
def attendance_close(session_id: int):
    with db() as c: c.execute("UPDATE attendance_sessions SET is_open=0 WHERE id=?", (session_id,))
    return redirect_with("/admin/attendance", msg="Đã đóng phiên điểm danh; điểm chuyên cần được cập nhật.")

@app.get("/admin/attendance/view/{session_id}", response_class=HTMLResponse, name="attendance_view")
def attendance_view(request: Request, session_id: int):
    with db() as c:
        sess = c.execute("SELECT * FROM attendance_sessions WHERE id=?", (session_id,)).fetchone()
        records = c.execute("""SELECT s.student_id code,s.name,ar.status,ar.checked_at FROM attendance_records ar JOIN students s ON s.id=ar.student_id WHERE ar.session_id=? ORDER BY s.stt""", (session_id,)).fetchall()
    if not sess: return PlainTextResponse("Không tìm thấy phiên", status_code=404)
    return render(request, "attendance_view.html", session=sess, records=records)

@app.get("/admin/attendance/{session_id}/qr.png", name="attendance_qr")
def attendance_qr(request: Request, session_id: int):
    with db() as c: sess = c.execute("SELECT token FROM attendance_sessions WHERE id=?", (session_id,)).fetchone()
    if not sess: return PlainTextResponse("Không tìm thấy phiên", status_code=404)
    url = str(request.base_url).rstrip("/") + f"/attendance/checkin/{sess['token']}"
    img = qrcode.make(url)
    buf = io.BytesIO(); img.save(buf, format="PNG"); buf.seek(0)
    return StreamingResponse(buf, media_type="image/png")

@app.get("/attendance/checkin/{token}", response_class=HTMLResponse, name="attendance_checkin")
def attendance_checkin(request: Request, token: str):
    with db() as c: sess = c.execute("SELECT * FROM attendance_sessions WHERE token=?", (token,)).fetchone()
    expired = True
    if sess:
        expired = (not sess["is_open"]) or datetime.now() > datetime.fromisoformat(sess["expires_at"])
    return render(request, "attendance_checkin.html", session=sess, expired=expired)

@app.post("/attendance/checkin/{token}", name="attendance_checkin_post")
def attendance_checkin_post(token: str, student_id: str = Form(...)):
    try:
        with db() as c:
            sess = c.execute("SELECT * FROM attendance_sessions WHERE token=?", (token,)).fetchone()
            if not sess or not sess["is_open"] or datetime.now() > datetime.fromisoformat(sess["expires_at"]):
                raise ValueError("Phiên điểm danh đã đóng hoặc hết hạn.")
            srow = c.execute("SELECT * FROM students WHERE student_id=?", (student_id.strip(),)).fetchone()
            if not srow: raise ValueError("Không tìm thấy MSSV/số thẻ.")
            c.execute("INSERT INTO attendance_records(session_id,student_id,status,checked_at) VALUES (?,?,?,?) ON CONFLICT(session_id,student_id) DO UPDATE SET status='Có mặt',checked_at=excluded.checked_at",
                      (sess["id"], srow["id"], "Có mặt", datetime.now().isoformat(timespec="seconds")))
        return redirect_with(f"/attendance/checkin/{token}", msg="Điểm danh thành công.")
    except Exception as exc:
        return redirect_with(f"/attendance/checkin/{token}", err=str(exc))

def assignment_context(group_id: int):
    with db() as c:
        g = c.execute("SELECT g.*,pd.data_code,pd.p,pd.v,pd.d,pd.t,pd.nam,pd.ngay FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id WHERE g.id=?", (group_id,)).fetchone()
        if not g:
            raise ValueError("Không tìm thấy nhóm")
        members = c.execute("SELECT s.* FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (group_id,)).fetchall()
        progress = c.execute("SELECT m.*,gp.status,gp.score,gp.feedback FROM milestones m LEFT JOIN group_progress gp ON gp.milestone_id=m.id AND gp.group_id=? ORDER BY m.position", (group_id,)).fetchall()
    if g["project_type"] not in PROJECT_INFO or not g["data_code"]:
        raise ValueError("Nhóm chưa được phân đề/số liệu.")
    return {"g": g, "members": members, "progress": progress, "info": PROJECT_INFO[g["project_type"]], "settings": get_settings()}


@app.get("/group/{group_id}/print", response_class=HTMLResponse, name="print_assignment")
def print_assignment(request: Request, group_id: int):
    try:
        ctx = assignment_context(group_id)
        return render(request, "assignment_print.html", **ctx)
    except Exception as exc:
        return redirect_with(f"/group/{group_id}", err=str(exc))


@app.get("/group/{group_id}/pdf", name="download_pdf")
def download_pdf(request: Request, group_id: int):
    # WeasyPrint needs native GTK/Pango libraries on Windows. If they are not
    # available, use the browser's built-in Print -> Save as PDF instead.
    if not WEASYPRINT_AVAILABLE:
        return RedirectResponse(url=f"/group/{group_id}/print", status_code=303)
    try:
        ctx = assignment_context(group_id)
        html = templates.get_template("assignment_pdf.html").render(
            request=request, PROJECT_INFO=PROJECT_INFO, format_num=format_num,
            load_image_url=load_image_url, system_image_url=system_image_url, base_dir=str(BASE_DIR), **ctx
        )
        out = GENERATED_DIR / f"{ctx['g']['group_code']}_DA{ctx['g']['project_type']}_{str(ctx['g']['data_code']).replace('.', '-')}.pdf"
        HTML(string=html, base_url=str(BASE_DIR)).write_pdf(out)
    except Exception as exc:
        return redirect_with(f"/group/{group_id}", err=str(exc))
    return FileResponse(out, media_type="application/pdf", filename=out.name)


if __name__ == "__main__":
    init_db()
    seed_demo_if_empty()
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", "5000")))
